#!/usr/bin/env python3
"""
inspect_resume.py — Structural ATS-risk scan for a resume file (.docx or .pdf).

This does NOT judge content/wording. It inspects the file's actual structure —
the same way an ATS parser would encounter it — and flags elements that commonly
cause resumes to be mis-parsed or dropped: tables, text boxes, headers/footers,
images/icons, multi-column layouts, non-standard fonts, missing standard section
headers, and (for PDFs) text that isn't actually extractable.

Usage:
    python inspect_resume.py <path-to-resume.docx-or-.pdf>

Output: JSON to stdout with `extracted_text`, `flags` (list of {severity, issue,
detail}), and a few raw structural stats for reference.
"""

import json
import re
import sys
from pathlib import Path

STANDARD_SECTION_WORDS = [
    "experience", "employment", "work history",
    "education",
    "skills", "expertise", "technical skills",
    "summary", "profile", "objective",
]


def inspect_docx(path):
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(path)
    flags = []

    # Tables
    n_tables = len(doc.tables)
    if n_tables > 0:
        flags.append({
            "severity": "high",
            "issue": f"{n_tables} table(s) detected in the document body",
            "detail": (
                "Many ATS parsers read tables cell-by-cell out of order, or skip them "
                "entirely — this is one of the most common causes of a resume being "
                "mis-parsed. If a table is used for layout (e.g. a two-column skills "
                "list, or a header row with contact info), convert it to plain "
                "paragraphs/line breaks instead."
            ),
        })

    # Headers / footers with real content
    header_text = ""
    footer_text = ""
    for section in doc.sections:
        header_text += "".join(p.text for p in section.header.paragraphs)
        footer_text += "".join(p.text for p in section.footer.paragraphs)
    if header_text.strip():
        flags.append({
            "severity": "high",
            "issue": "Content found in the document header",
            "detail": (
                f"Header text: {header_text.strip()[:200]!r}. Many ATS systems "
                "ignore headers/footers entirely — if contact info (name, phone, "
                "email) lives only in the header, it may never be captured. Move "
                "it into the main document body instead."
            ),
        })
    if footer_text.strip():
        flags.append({
            "severity": "medium",
            "issue": "Content found in the document footer",
            "detail": (
                f"Footer text: {footer_text.strip()[:200]!r}. Same risk as header "
                "content — footers are frequently skipped by parsers."
            ),
        })

    # Text boxes / floating shapes (stored as w:pict or mc:AlternateContent in the XML)
    xml = doc.element.xml
    textbox_count = xml.count("<w:pict") + xml.count("<v:textbox")
    if textbox_count > 0:
        flags.append({
            "severity": "high",
            "issue": f"~{textbox_count} text box(es) or floating drawing object(s) detected",
            "detail": (
                "Text inside text boxes or floating shapes is frequently invisible "
                "to ATS parsers, even though it displays fine visually. Any content "
                "in a text box should be moved into the normal document flow."
            ),
        })

    # Inline images (logos, icons used as bullet/section markers)
    n_images = len(doc.inline_shapes)
    if n_images > 0:
        flags.append({
            "severity": "low",
            "issue": f"{n_images} inline image(s)/icon(s) detected",
            "detail": (
                "Images and icons (including icon fonts used for contact info like "
                "phone/email/LinkedIn glyphs) are invisible to text-based ATS "
                "parsing. Any information conveyed only through an icon (e.g. a "
                "phone icon next to a number with no 'Phone:' label) should also "
                "exist as plain text nearby."
            ),
        })

    # Fonts used
    fonts = set()
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.name:
                fonts.add(run.font.name)
    unusual_fonts = {f for f in fonts if f and f.lower() not in {
        "calibri", "arial", "times new roman", "georgia", "helvetica",
        "cambria", "garamond", "verdana", "tahoma", "book antiqua",
    }}
    if unusual_fonts:
        flags.append({
            "severity": "low",
            "issue": f"Non-standard font(s) in use: {', '.join(sorted(unusual_fonts))}",
            "detail": (
                "Decorative or uncommon fonts can render as garbled characters or "
                "fail to embed correctly when an ATS converts the file to plain "
                "text. Standard, widely-installed fonts parse most reliably."
            ),
        })

    # Extract full text for content matching downstream
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            full_text += "\n" + " | ".join(cell.text for cell in row.cells)

    _check_section_headers(full_text, flags)

    return full_text, flags, {"n_tables": n_tables, "n_images": n_images}


def inspect_pdf(path):
    import pdfplumber

    flags = []
    full_text = ""
    total_images = 0
    font_names = set()
    column_risk_pages = 0

    with pdfplumber.open(path) as pdf:
        n_pages = len(pdf.pages)
        for page in pdf.pages:
            text = page.extract_text() or ""
            full_text += text + "\n"
            total_images += len(page.images)
            for ch in page.chars[:500]:  # sample for font names, avoid huge scans
                if ch.get("fontname"):
                    font_names.add(ch["fontname"])

            # Rough multi-column heuristic: cluster word x0 positions into left/right
            # halves of the page; if a large fraction of words sit in a narrow band
            # on both halves (rather than spanning full width), it's likely a
            # multi-column layout, which ATS parsers often read left-to-right across
            # columns and scramble.
            words = page.extract_words()
            if words:
                page_width = page.width
                left = [w for w in words if w["x0"] < page_width * 0.45]
                right = [w for w in words if w["x0"] > page_width * 0.55]
                spanning = [w for w in words if page_width * 0.4 < w["x0"] < page_width * 0.6]
                if len(left) > 5 and len(right) > 5 and len(spanning) < 0.1 * len(words):
                    column_risk_pages += 1

    if not full_text.strip():
        flags.append({
            "severity": "critical",
            "issue": "No extractable text found in the PDF",
            "detail": (
                "This PDF appears to contain no machine-readable text at all — it's "
                "likely a scanned image or an exported image-only PDF. Standard ATS "
                "parsers cannot read this at all without OCR, and most do not OCR "
                "uploaded resumes. This is disqualifying for ATS parsing until fixed "
                "— export a true text-based PDF or DOCX instead."
            ),
        })

    if total_images > 0:
        flags.append({
            "severity": "low",
            "issue": f"{total_images} embedded image(s)/icon(s) detected across the document",
            "detail": (
                "Same risk as in Word documents — any information conveyed only "
                "through an image or icon glyph won't be captured as text."
            ),
        })

    if column_risk_pages > 0:
        flags.append({
            "severity": "high",
            "issue": f"Likely multi-column layout detected on {column_risk_pages} page(s)",
            "detail": (
                "Multi-column resume layouts (e.g. a narrow sidebar for skills/"
                "contact info next to a wider experience column) are commonly read "
                "left-to-right straight across both columns by ATS text extraction, "
                "scrambling the order — a bullet from the sidebar can end up "
                "spliced into the middle of a job description. A single-column "
                "layout is the safest choice."
            ),
        })

    _check_section_headers(full_text, flags)

    return full_text, flags, {"n_pages": n_pages, "n_images": total_images}


def _check_section_headers(text, flags):
    lower = text.lower()
    missing = [w for w in ["experience", "education", "skills"]
               if not any(w2 in lower for w2 in
                          {"experience": ["experience", "employment", "work history"],
                           "education": ["education"],
                           "skills": ["skills", "expertise"]}[w])]
    if missing:
        flags.append({
            "severity": "medium",
            "issue": f"Missing conventional section header(s): {', '.join(missing)}",
            "detail": (
                "ATS parsers typically segment a resume by matching standard "
                "section headers (Experience, Education, Skills, etc.). If a "
                "section uses a creative or non-standard label instead (e.g. "
                "'My Journey' instead of 'Experience'), the parser may fail to "
                "recognize and correctly bucket that content."
            ),
        })

    # Unusual bullet glyphs that sometimes fail to convert cleanly to plain text
    unusual_bullets = set(re.findall(r"[\u2726\u27a4\u2756\uf0b7\u25c6\u2765]", text))
    if unusual_bullets:
        flags.append({
            "severity": "low",
            "issue": "Non-standard bullet character(s) detected",
            "detail": (
                "Decorative bullet glyphs (as opposed to a standard '•' or '-') can "
                "sometimes render as garbled characters (e.g. a stray letter or "
                "box symbol) after ATS text extraction."
            ),
        })


def main():
    if len(sys.argv) != 2:
        print("Usage: python inspect_resume.py <path-to-resume.docx-or-.pdf>", file=sys.stderr)
        sys.exit(1)

    path = Path(sys.argv[1])
    if not path.exists():
        print(json.dumps({"error": f"File not found: {path}"}))
        sys.exit(1)

    ext = path.suffix.lower()
    if ext == ".docx":
        text, flags, stats = inspect_docx(str(path))
    elif ext == ".pdf":
        text, flags, stats = inspect_pdf(str(path))
    else:
        print(json.dumps({"error": f"Unsupported file type: {ext}. Use .docx or .pdf."}))
        sys.exit(1)

    severity_order = {"critical": 0, "high": 1, "medium": 2, "low": 3}
    flags.sort(key=lambda f: severity_order.get(f["severity"], 9))

    print(json.dumps({
        "file": str(path),
        "stats": stats,
        "flags": flags,
        "extracted_text": text,
    }, indent=2))


if __name__ == "__main__":
    main()
